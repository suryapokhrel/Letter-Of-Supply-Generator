from django.shortcuts import render
from django.http import HttpResponse
from django.conf import settings
from .forms import LOSForm
import os
from docx.api import Document

from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml import OxmlElement


def generate_letter_of_supply(request):
    if request.method == 'POST':
        form = LOSForm(request.POST, request.FILES)
        if form.is_valid():
            vendorName = form.cleaned_data['vendorName']
            vendorPOCName = form.cleaned_data['vendorPOCName']
            vendorPOCTitle = form.cleaned_data['vendorPOCTitle']
            vendorAddress = form.cleaned_data['vendorAddress']
            supplierName = form.cleaned_data['supplierName']
            supplierPOCName = form.cleaned_data['supplierPOCName']
            supplierPOCTitle = form.cleaned_data['supplierPOCTitle']
            supplierBrands = form.cleaned_data['supplierBrands']
            losType = form.cleaned_data['losType']
            contractNumber = form.cleaned_data['contractNumber']
            uploaded_file = request.FILES['letterhead']
            checkboxOptions = form.cleaned_data.get('specialSINs', [])
            checkboxValues = set([checkbox_option for checkbox_option in checkboxOptions])
            radioOption = form.cleaned_data.get('losType', '')
            radioValues = set(['NEW'] if radioOption == 'NEW CONTRACT' else ['UPDATE'] if radioOption == 'UPDATE CONTRACT' else [])
            
            doc = Document('temp.docx')
            paragraph = doc.paragraphs[0].insert_paragraph_before()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img = paragraph.add_run()
            img.add_picture(uploaded_file, width=Inches(1.0)) 
            print(form.cleaned_data.get('specialSINs', []))
            data = {
                '[Vendor Name]': vendorName,
                '[Vendor POC Name]': vendorPOCName,
                '[Vendor POC Title]': vendorPOCTitle,
                '[Vendor Address]': vendorAddress,
                '[Supplier Name]': supplierName,
                '[Supplier POC Name]': supplierPOCName,
                '[Supplier POC Title]': supplierPOCTitle,
                '[Supplier Brands]': supplierBrands,
                '[LOS TYPE]': losType,
                '[Contract Number]': contractNumber,
                '[NEW]': 'X_[NEW]_X' if losType == 'NEW' else '__[NEW]__',
                '[UPDATE]': "X_[UPDATE]_X" if losType == 'UPDATE' else '__[UPDATE]__',
                '[IT]': '_X_[IT]_X_' if '[IT]' in form.cleaned_data.get('specialSINs', []) else '__[IT]__',
                '[OFFICE]': '_X_[OFFICE]_X_' if '[OFFICE]' in form.cleaned_data.get('specialSINs', []) else '__[OFFICE]__',
            }
                                

            for paragraph in doc.paragraphs:
                for key, value in data.items():
                    if key in paragraph.text:
                        inline = paragraph.runs
                        for i in range(len(inline)):
                            if key in inline[i].text:
                                inline[i].text = inline[i].text.replace(key, value)
                                # if inline[i].italic:
                                #     font = inline[i].font
                                #     font.italic = None
                                #     italic_element = OxmlElement('w:i')
                                #     inline[i].element.rPr.append(italic_element)
                        for run in inline:
                            run.font.italic = None
                    
            output_file_path = os.path.join(settings.MEDIA_ROOT, 'modified_letter_of_supply.docx')
            doc.save(output_file_path)
            with open(output_file_path, 'rb') as modified_file:
                response = HttpResponse(modified_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = 'attachment; filename=modified_letter_of_supply.docx'
                return response
    else:
        form = LOSForm()

    return render(request, 'home.html', {'form': form})
